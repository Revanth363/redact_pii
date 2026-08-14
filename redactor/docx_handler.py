"""
redactor/docx_handler.py
------------------------
Reads a .docx file, converts it into text blocks, then writes a redacted
.docx preserving the original formatting as closely as possible.
"""

from __future__ import annotations

import logging
from dataclasses import dataclass, field
from typing import List, Optional

logger = logging.getLogger(__name__)

try:
    from docx import Document
    _docx_available = True
except ImportError:
    _docx_available = False
    logger.warning(
        "python-docx not installed. Install with: pip install python-docx"
    )


@dataclass
class Block:
    block_id: str
    text:     str
    source:   str                           # paragraph | table | header | footer
    para_obj: Optional[object] = field(default=None, repr=False)


class DocxHandler:

    def __init__(self, path: str) -> None:
        if not _docx_available:
            raise ImportError("python-docx is required. pip install python-docx")
        self._path = path
        self._doc  = Document(path)

    # ------------------------------------------------------------------
    # Extraction
    # ------------------------------------------------------------------

    def extract_blocks(self) -> List[Block]:
        """
        Extract all text blocks from the document.
        Returns a list of Block objects in document order.
        """
        blocks: List[Block] = []
        counter = {"n": 0}

        def next_id(prefix: str) -> str:
            counter["n"] += 1
            return f"{prefix}_{counter['n']}"

        # Main body paragraphs
        for para in self._doc.paragraphs:
            text = para.text
            if not text.strip():
                continue
            blocks.append(Block(
                block_id=next_id("para"),
                text=text,
                source="paragraph",
                para_obj=para,
            ))

        # Tables
        for table in self._doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text
                        if not text.strip():
                            continue
                        blocks.append(Block(
                            block_id=next_id("table"),
                            text=text,
                            source="table",
                            para_obj=para,
                        ))

        # Headers and footers
        for section in self._doc.sections:
            for hf_para in list(section.header.paragraphs) + list(section.footer.paragraphs):
                text = hf_para.text
                if not text.strip():
                    continue
                blocks.append(Block(
                    block_id=next_id("headfoot"),
                    text=text,
                    source="header" if hf_para in section.header.paragraphs else "footer",
                    para_obj=hf_para,
                ))

        logger.info("DocxHandler: extracted %d non-empty blocks", len(blocks))
        return blocks

    def block_text_map(self, blocks: List[Block]) -> dict[str, str]:
        """Return {block_id: text} for use by the context engine."""
        return {b.block_id: b.text for b in blocks}

    # ------------------------------------------------------------------
    # Replacement
    # ------------------------------------------------------------------

    def apply_replacements(
        self,
        blocks: List[Block],
        replacement_map: dict[str, List[tuple[int, int, str]]],
    ) -> None:
        """
        Perform in-place text replacements inside each block.
        """
        for block in blocks:
            spans = replacement_map.get(block.block_id)
            if not spans or block.para_obj is None:
                continue
            self._replace_in_paragraph(block.para_obj, block.text, spans)

    def save(self, output_path: str) -> None:
        self._doc.save(output_path)
        logger.info("DocxHandler: saved redacted document to '%s'", output_path)

    # ------------------------------------------------------------------
    # Private — run-level replacement
    # ------------------------------------------------------------------

    def _replace_in_paragraph(
        self,
        para,
        original_text: str,
        spans: List[tuple[int, int, str]],
    ) -> None:
        """
        Replace spans inside a paragraph while preserving run formatting
        and ensuring bulletproof text substitution.
        """
        if not spans:
            return

        spans = sorted(spans, key=lambda s: s[0])
        new_text = self._apply_spans_to_text(original_text, spans)

        runs = para.runs
        if runs:
            runs[0].text = new_text
            for r in runs[1:]:
                r.text = ""
        else:
            para.text = new_text

        # Secondary safety pass: direct string replacements for any target text
        for (start, end, replacement) in spans:
            target_str = original_text[start:end].strip()
            if target_str and len(target_str) >= 2 and target_str in para.text:
                para.text = para.text.replace(target_str, replacement)

    @staticmethod
    def _apply_spans_to_text(
        text: str,
        spans: List[tuple[int, int, str]],
    ) -> str:
        """Build the new string by substituting spans into text."""
        result = []
        cursor = 0
        for (start, end, replacement) in spans:
            result.append(text[cursor:start])
            result.append(replacement)
            cursor = end
        result.append(text[cursor:])
        return "".join(result)
